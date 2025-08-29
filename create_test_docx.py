#!/usr/bin/env python3
import sys
try:
    from docx import Document
    from docx.shared import RGBColor
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import RGBColor

# Create a new document
doc = Document()

# Add title
doc.add_heading('生物知识测试', 0)

# Add content with highlights
p = doc.add_paragraph('这是一个测试文档，用于验证LLM集成是否正常工作。\n\n')

# Add highlighted content
p = doc.add_paragraph('重要知识点：\n')
p = doc.add_paragraph()
run = p.add_run('光合作用')
run.font.highlight_color = 7  # Yellow highlight
p.add_run('是植物将光能转化为化学能的过程。\n')

p = doc.add_paragraph()
run = p.add_run('线粒体')
run.font.highlight_color = 7  # Yellow highlight
p.add_run('被称为细胞的"动力工厂"。\n')

p = doc.add_paragraph()
run = p.add_run('DNA')
run.font.highlight_color = 7  # Yellow highlight
p.add_run('携带遗传信息。\n')

p = doc.add_paragraph('\n请根据以上黄色高亮内容生成选择题和填空题。')

# Save the document
doc.save('test-quiz.docx')
print("Created test-quiz.docx successfully!")